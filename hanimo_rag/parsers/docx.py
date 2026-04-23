"""DOCX parser using python-docx."""
from __future__ import annotations

from hanimo_rag.parsers.base import ParsedDocument, ParserBase

# Heading level → markdown prefix
_HEADING_PREFIX: dict[str, str] = {
    "Heading 1": "# ",
    "Heading 2": "## ",
    "Heading 3": "### ",
    "Heading 4": "#### ",
    "Heading 5": "##### ",
    "Heading 6": "###### ",
}


def _table_to_markdown(table) -> str:
    """Convert a python-docx Table to a markdown table string."""
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(rows)


class DocxParser(ParserBase):
    """Parse DOCX files preserving heading structure and tables."""

    def supported_mime_types(self) -> list[str]:
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ]

    def parse(self, file_path: str) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError as e:
            return ParsedDocument(
                text="",
                metadata={"error": f"Missing dependency: {e}"},
            )

        try:
            doc = Document(file_path)
        except Exception as exc:
            return ParsedDocument(text="", metadata={"error": str(exc)})

        # --- Metadata from core properties ---
        metadata: dict = {}
        try:
            props = doc.core_properties
            metadata["title"] = props.title or ""
            metadata["author"] = props.author or ""
            metadata["created"] = str(props.created) if props.created else ""
        except Exception:
            pass

        # --- Build content: paragraphs + tables in document order ---
        parts: list[str] = []

        # Iterate over body children to preserve order of paragraphs and tables
        body = doc.element.body
        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":
                # It's a paragraph
                from docx.text.paragraph import Paragraph
                para = Paragraph(child, doc)
                style_name = para.style.name if para.style else ""
                text = para.text
                if not text.strip():
                    continue
                prefix = _HEADING_PREFIX.get(style_name, "")
                parts.append(f"{prefix}{text}")

            elif tag == "tbl":
                # It's a table
                from docx.table import Table
                table = Table(child, doc)
                parts.append(_table_to_markdown(table))

        full_text = "\n\n".join(parts)
        # Pages = one page (DOCX doesn't have native page breaks we can easily split on)
        pages = [full_text] if full_text else []

        return ParsedDocument(text=full_text, metadata=metadata, pages=pages)
